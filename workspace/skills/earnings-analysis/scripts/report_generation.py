# scripts/report_generation.py
"""
报告生成脚本：整合所有数据和分析，生成最终报告。
"""
from docx import Document
from docx.shared import Inches
from matplotlib import pyplot as plt
import pandas as pd

def create_earnings_update_report(company_name, ticker, quarter, analysis_data, web_data, output_path):
    """
    创建财报更新报告。

    Args:
        company_name (str): 公司名称
        ticker (str): 股票代码
        quarter (str): 季度
        analysis_data (dict): 来自财务分析的数据
        web_data (dict): 来自网络搜索的数据
        output_path (str): 输出文件路径
    """
    doc = Document()

    # 添加标题页
    add_title_page(doc, company_name, ticker, quarter)

    # 添加摘要页
    add_summary_section(doc, company_name, quarter, analysis_data, web_data)

    # 添加详细分析
    add_detailed_analysis(doc, analysis_data)

    # 添加图表
    add_charts(doc, analysis_data)

    # 添加来源和参考
    add_sources_section(doc, web_data)

    # 保存文档
    doc.save(output_path)

def add_title_page(doc, company_name, ticker, quarter):
    """添加报告标题页"""
    doc.add_heading(f"{company_name} ({ticker}) {quarter} EARNINGS UPDATE", 0)
    doc.add_paragraph(f"Date: {pd.Timestamp.now().strftime('%B %d, %Y')}")

def add_summary_section(doc, company_name, quarter, analysis_data, web_data):
    """添加摘要部分"""
    doc.add_heading("EARNINGS SUMMARY", level=1)

    # 创建摘要框
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # 填充摘要信息
    metrics = analysis_data['metrics']

    table.rows[0].cells[0].text = f"{quarter} RESULTS:"
    table.rows[0].cells[1].text = "BEAT"  # 需要实际逻辑判断

    table.rows[1].cells[0].text = "Revenue ($M)"
    table.rows[1].cells[1].text = f"{metrics['revenue']:,.0f}"

    table.rows[2].cells[0].text = "EPS (Adj) ($)"
    table.rows[2].cells[1].text = f"{metrics['eps']:,.2f}"

    table.rows[3].cells[0].text = "Key Takeaways:"
    table.rows[3].cells[1].text = generate_key_takeaways(analysis_data, web_data)

    table.rows[4].cells[0].text = "Rating Impact:"
    table.rows[4].cells[1].text = "MAINTAIN OUTPERFORM"  # 需要实际逻辑判断

def add_detailed_analysis(doc, analysis_data):
    """添加详细分析部分"""
    doc.add_page_break()
    doc.add_heading("DETAILED RESULTS ANALYSIS", level=1)

    # 添加收入分析
    add_revenue_analysis(doc, analysis_data)

    # 添加盈利能力分析
    add_profitability_analysis(doc, analysis_data)

    # 添加现金流分析
    add_cash_flow_analysis(doc, analysis_data)

def add_charts(doc, analysis_data):
    """添加图表"""
    doc.add_page_break()
    doc.add_heading("KEY TRENDS", level=1)

    # 收入趋势图
    revenue_chart = create_trend_chart(analysis_data['revenue_trend'], "Revenue Trend")
    doc.add_picture(revenue_chart, width=Inches(6))

    # 盈利能力趋势图
    margin_chart = create_trend_chart(analysis_data['margin_trend'], "Margin Trends")
    doc.add_picture(margin_chart, width=Inches(6))

def add_sources_section(doc, web_data):
    """添加来源部分"""
    doc.add_page_break()
    doc.add_heading("SOURCES & REFERENCES", level=1)

    # 添加数据来源
    p = doc.add_paragraph()
    p.add_run("Financial Data: ").bold = True
    p.add_run("Company 10-Q and 10-K filings")

    # 添加网络搜索来源
    if web_data and 'earnings_release' in web_data:
        p = doc.add_paragraph()
        p.add_run("Earnings Release: ").bold = True
        p.add_run(web_data['earnings_release']['title'])
        p.add_run(" [Link: ").add_hyperlink(web_data['earnings_release']['url'])
        p.add_run("]")

def generate_key_takeaways(analysis_data, web_data):
    """生成关键要点"""
    # 这里应该有更复杂的逻辑生成实际要点
    return "1. Revenue grew X% YoY, driven by [segment]\n2. Operating margin expanded Xbps\n3. Free cash flow generation remained strong"

def create_trend_chart(data, title):
    """创建趋势图"""
    fig, ax = plt.subplots(figsize=(10, 6))

    if isinstance(data, pd.DataFrame):
        for column in data.columns:
            ax.plot(data.index, data[column], marker='o', label=column)
        ax.legend()
    else:
        ax.plot(data['labels'], data['values'], marker='o')

    ax.set_title(title)
    ax.set_xlabel('Quarter')
    ax.set_ylabel('Value')
    ax.grid(True)

    plt.tight_layout()
    return fig
